import pandas as pd
import numpy as np
from docx import Document
from datetime import datetime

# ==============================
# Simulación Grid Search
# ==============================

np.random.seed(42)

results = []

for rsi in range(10, 30, 5):
    for adx in range(5, 20, 5):
        for atr in range(5, 20, 5):
            profit = np.random.uniform(0, 100)
            drawdown = np.random.uniform(0, 50)
            sharpe = np.random.uniform(0, 3)

            results.append({
                "RSI": rsi,
                "ADX": adx,
                "ATR": atr,
                "Profit %": round(profit, 2),
                "Max Drawdown %": round(drawdown, 2),
                "Sharpe Ratio": round(sharpe, 2)
            })

df = pd.DataFrame(results)

# Mejor combinación por Profit
best = df.sort_values("Profit %", ascending=False).iloc[0]

# ==============================
# Crear Word
# ==============================

document = Document()
document.add_heading("Reporte Automático de Grid Search", level=1)
document.add_paragraph(f"Fecha: {datetime.now()}")

document.add_heading("Mejor Configuración", level=2)
document.add_paragraph(f"RSI: {best['RSI']}")
document.add_paragraph(f"ADX: {best['ADX']}")
document.add_paragraph(f"ATR: {best['ATR']}")
document.add_paragraph(f"Profit %: {best['Profit %']}")
document.add_paragraph(f"Max Drawdown %: {best['Max Drawdown %']}")
document.add_paragraph(f"Sharpe Ratio: {best['Sharpe Ratio']}")

document.add_heading("Tabla completa de resultados", level=2)

table = document.add_table(rows=1, cols=len(df.columns))
hdr_cells = table.rows[0].cells

for i, column in enumerate(df.columns):
    hdr_cells[i].text = column

for _, row in df.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)

document.save("grid_report.docx")

print("Reporte generado correctamente.")
